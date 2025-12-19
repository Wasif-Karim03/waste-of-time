"""Resume parsing functionality to extract skills, keywords, and information."""

import re
import io
from typing import Dict, List, Set
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Common technical skills keywords
TECHNICAL_SKILLS = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'go', 'golang',
    'rust', 'ruby', 'php', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl',
    'shell', 'bash', 'powershell', 'sql', 'html', 'css', 'sass', 'scss',
    
    # Frameworks & Libraries
    'react', 'angular', 'vue', 'node.js', 'nodejs', 'express', 'django', 'flask',
    'fastapi', 'spring', 'spring boot', 'laravel', 'rails', 'ruby on rails',
    'asp.net', '.net', 'jquery', 'bootstrap', 'tailwind', 'next.js', 'nuxt',
    
    # Databases
    'mysql', 'postgresql', 'postgres', 'mongodb', 'redis', 'elasticsearch',
    'cassandra', 'dynamodb', 'oracle', 'sqlite', 'mariadb', 'neo4j',
    
    # Cloud & DevOps
    'aws', 'amazon web services', 'azure', 'gcp', 'google cloud', 'docker',
    'kubernetes', 'k8s', 'terraform', 'ansible', 'jenkins', 'ci/cd', 'github actions',
    'gitlab ci', 'circleci', 'travis ci', 'terraform', 'cloudformation',
    
    # Tools & Technologies
    'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack',
    'agile', 'scrum', 'kanban', 'rest api', 'graphql', 'microservices',
    'linux', 'unix', 'windows', 'macos', 'nginx', 'apache',
    
    # Data & AI/ML
    'machine learning', 'ml', 'deep learning', 'neural networks', 'tensorflow',
    'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'data science',
    'data analysis', 'big data', 'hadoop', 'spark', 'kafka', 'airflow',
    
    # Mobile
    'ios', 'android', 'react native', 'flutter', 'xamarin', 'ionic',
    
    # Other
    'blockchain', 'ethereum', 'solidity', 'web3', 'cybersecurity', 'security',
    'testing', 'unit testing', 'integration testing', 'tdd', 'bdd'
}

# Soft skills keywords
SOFT_SKILLS = {
    'leadership', 'communication', 'teamwork', 'collaboration', 'problem solving',
    'critical thinking', 'analytical', 'creative', 'innovative', 'adaptable',
    'flexible', 'time management', 'project management', 'agile', 'scrum master',
    'mentoring', 'coaching', 'presentation', 'public speaking', 'negotiation',
    'customer service', 'stakeholder management', 'cross-functional'
}

# Education keywords
EDUCATION_KEYWORDS = {
    'bachelor', 'master', 'phd', 'doctorate', 'degree', 'university', 'college',
    'bs', 'ba', 'ms', 'ma', 'mba', 'ph.d', 'computer science', 'engineering',
    'information technology', 'it', 'software engineering', 'data science'
}

# Experience keywords
EXPERIENCE_KEYWORDS = {
    'years of experience', 'yoe', 'experience', 'worked', 'developed', 'designed',
    'implemented', 'managed', 'led', 'architected', 'built', 'created'
}


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file."""
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
    
    text = ""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")
    
    return text


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
    
    try:
        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")
    
    return text


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file."""
    try:
        return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        raise ValueError(f"Error reading TXT: {str(e)}")


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from uploaded file based on file extension."""
    file_ext = Path(filename).suffix.lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_content)
    elif file_ext in ['.doc', '.docx']:
        return extract_text_from_docx(file_content)
    elif file_ext == '.txt':
        return extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")


def normalize_text(text: str) -> str:
    """Normalize text for better matching."""
    # Convert to lowercase
    text = text.lower()
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    return text


def find_skills_in_text(text: str, skill_set: Set[str]) -> List[str]:
    """Find skills from a set in the text."""
    normalized_text = normalize_text(text)
    found_skills = []
    
    for skill in skill_set:
        # Check for exact word boundaries
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, normalized_text, re.IGNORECASE):
            found_skills.append(skill.title())
    
    return sorted(set(found_skills))


def extract_education(text: str) -> List[str]:
    """Extract education information."""
    education = []
    normalized_text = normalize_text(text)
    
    # Look for degree patterns
    degree_patterns = [
        r'\b(bachelor|b\.?s\.?|b\.?a\.?)\s+(?:of\s+)?(?:science|arts|engineering)?',
        r'\b(master|m\.?s\.?|m\.?a\.?|m\.?b\.?a\.?)\s+(?:of\s+)?(?:science|arts|business|engineering)?',
        r'\b(ph\.?d\.?|doctorate|doctor\s+of\s+philosophy)',
        r'\b(associate|a\.?s\.?|a\.?a\.?)\s+(?:of\s+)?(?:science|arts)?'
    ]
    
    for pattern in degree_patterns:
        matches = re.findall(pattern, normalized_text, re.IGNORECASE)
        if matches:
            education.extend([m[0] if isinstance(m, tuple) else m for m in matches])
    
    # Look for field of study
    field_patterns = [
        r'(?:in|of)\s+(computer\s+science|software\s+engineering|information\s+technology|it|data\s+science|engineering)',
        r'(computer\s+science|software\s+engineering|information\s+technology|it|data\s+science)'
    ]
    
    for pattern in field_patterns:
        matches = re.findall(pattern, normalized_text, re.IGNORECASE)
        education.extend(matches)
    
    return sorted(set(education))


def extract_experience_years(text: str) -> str:
    """Extract years of experience."""
    normalized_text = normalize_text(text)
    
    # Look for patterns like "5 years", "10+ years", etc.
    patterns = [
        r'(\d+)\+?\s*years?\s+(?:of\s+)?(?:experience|exp)',
        r'(?:experience|exp)[:\s]+(\d+)\+?\s*years?',
        r'(\d+)\+?\s*y\.?o\.?e\.?'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, normalized_text, re.IGNORECASE)
        if match:
            return f"{match.group(1)} years"
    
    return "Not specified"


def extract_keywords(text: str) -> List[str]:
    """Extract important keywords from resume."""
    normalized_text = normalize_text(text)
    
    # Extract capitalized phrases (likely important terms)
    keywords = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
    
    # Filter out common words and keep meaningful terms
    common_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
    keywords = [kw for kw in keywords if kw.lower() not in common_words and len(kw) > 3]
    
    # Limit to top 20 unique keywords
    return sorted(set(keywords))[:20]


def parse_resume(file_content: bytes, filename: str) -> Dict:
    """
    Parse resume and extract all relevant information.
    
    Args:
        file_content: Binary content of the resume file
        filename: Name of the file
    
    Returns:
        Dictionary containing extracted information:
        - technical_skills: List of technical skills found
        - soft_skills: List of soft skills found
        - education: List of education information
        - experience_years: Years of experience
        - keywords: List of important keywords
        - raw_text: Extracted text from resume
    """
    # Extract text from file
    text = extract_text_from_file(file_content, filename)
    
    # Extract information
    technical_skills = find_skills_in_text(text, TECHNICAL_SKILLS)
    soft_skills = find_skills_in_text(text, SOFT_SKILLS)
    education = extract_education(text)
    experience_years = extract_experience_years(text)
    keywords = extract_keywords(text)
    
    return {
        'technical_skills': technical_skills,
        'soft_skills': soft_skills,
        'education': education,
        'experience_years': experience_years,
        'keywords': keywords,
        'raw_text': text[:1000]  # First 1000 chars for preview
    }
